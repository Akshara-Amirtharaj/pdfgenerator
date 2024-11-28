import streamlit as st
from docx import Document
import os
# Function to edit the Word template dynamically
def edit_word_template(template_path, output_path, name, designation, contact, email, location, selected_services):
    try:
        doc = Document(template_path)
        # Replace placeholders in the general paragraphs
        for para in doc.paragraphs:
            if "<<Client Name>>" in para.text:
                para.text = para.text.replace("<<Client Name>>", name)
            if "<<Client Designation>>" in para.text:
                para.text = para.text.replace("<<Client Designation>>", designation)
            if "<<Client Contact>>" in para.text:
                para.text = para.text.replace("<<Client Contact>>", contact)
            if "<<Client Email>>" in para.text:
                para.text = para.text.replace("<<Client Email>>", email)
            if "<<Client Location>>" in para.text:
                para.text = para.text.replace("<<Client Location>>", location)
                
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "<<Client Name>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Name>>", name)
                    if "<<Client Designation>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Designation>>", designation)
                    if "<<Client Contact>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Contact>>", contact)
                    if "<<Client Email>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Email>>", email)
                    if "<<Client Location>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Location>>", location)
        # Process all tables
        spoc_table_found = False  # Flag to indicate if the SPOC table is found
        for table_idx, table in enumerate(doc.tables):
            # Check for the SPOC table by searching for the text "Supporting SPOC Details"
            if not spoc_table_found:  # Look for the SPOC identifier
                for para in doc.paragraphs:
                    if "Supporting SPOC Details" in para.text:
                        spoc_table_found = True
                        break
            if spoc_table_found and table_idx == 0:  # Assuming SPOC table is the first table after the identifier
                # Update placeholders in the SPOC table
                for row in table.rows:
                    if "Project Sponsor/Client’s Detail" in row.cells[0].text:
                        row.cells[1].text = name
                        row.cells[2].text = designation
                        row.cells[3].text = contact
                        row.cells[4].text = email
                spoc_table_found = False  # Reset the flag after processing the table
            else:
                # Filter rows based on selected services for other tables
                for row in table.rows[1:]:  # Skip the header row
                    service_name = row.cells[0].text.strip()
                    if service_name not in selected_services:
                        row._element.getparent().remove(row._element)
        # Save the updated document
        doc.save(output_path)
        print(f"Word document updated and saved at: {output_path}")
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")
# Updated convert_to_pdf function
def convert_to_pdf(doc_path, pdf_path):
    word = None
    try:
        import comtypes.client
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        print(f"Converted to PDF and saved at: {pdf_path}")
    except Exception as e:
        raise Exception(f"Error converting Word to PDF: {e}")
    finally:
        if word:
            word.Quit()
# Streamlit App


st.title("Client-Specific PDF Generator")
# Input fields
name = st.text_input("Name")
designation = st.text_input("Designation")
contact = st.text_input("Contact Number")
email = st.text_input("Email ID")
location = st.selectbox("Location", ["India", "ROW"])

# List of all available services (ensure this matches your template)
services = [
    "Landing page website (design + development)",
    "AI Automations (6 Scenarios)",
    "WhatsApp Automation + WhatsApp Cloud Business Account Setup",
    "CRM Setup",
    "Email Marketing Setup",
    "Make/Zapier Automation Setup",
    "Firefly Meeting Automation",
    "Marketing Strategy",
    "Social Media Channels",
    "Creatives (10 Per Month)",
    "Creatives (20 Per Month)",
    "Creatives (30 Per Month)",
    "Reels (10 Reels)",
    "Meta Ad Account Setup & Pages Setup",
    "Paid Ads (Lead Generation)",
    "Monthly Maintenance & Reporting",
    "AI Chatbot",
    "PDF Generation Automations",
    "AI Generated Social Media Content & Calendar",
    "Custom AI Models & Agents"
]

# Checkbox to select all services
select_all = st.checkbox("Select All Services")
if select_all:
    selected_services = services
else:
    selected_services = st.multiselect("Select Services", services)

# Define paths
base_dir = os.path.abspath(os.path.dirname(__file__))
template_path = os.path.join(base_dir, "DM & Automations Services Pricing - Andrew.docx")
word_output_path = os.path.join(base_dir, "Customized_Pricing.docx")
pdf_output_path = os.path.join(base_dir, "Customized_Pricing.pdf")

if st.button("Generate PDF"):
    if not all([name, designation, contact, email, location]) or not selected_services:
        st.error("All fields and at least one service must be selected!")
    else:
        try:
            edit_word_template(
                template_path, word_output_path, name, designation, contact, email, location, selected_services
            )
            convert_to_pdf(word_output_path, pdf_output_path)
            st.success("PDF generated successfully!")
            with open(pdf_output_path, "rb") as file:
                st.download_button("Download PDF", file, file_name="Customized_Pricing.pdf")
        except Exception as e:
            st.error(f"An error occurred: {e}")